import io
from pathlib import Path
import pymupdf  # PyMuPDF
from docx import Document
from app.services.rag.text_cleaner import clean_text

def extract_pdf_bytes(file_bytes: bytes) -> str:
    """Extract text from a PDF file stream."""
    text = []
    with pymupdf.open(stream=file_bytes, filetype="pdf") as pdf:
        for page in pdf:
            page_text = page.get_text()
            if page_text:
                text.append(page_text)
    return "\n".join(text).strip()

def extract_docx_bytes(file_bytes: bytes) -> str:
    """Extract text from a DOCX file stream, including tables."""
    document = Document(io.BytesIO(file_bytes))
    content = []
    
    for element in document.element.body:
        if element.tag.endswith('p'):
            for p in document.paragraphs:
                if p._element == element and p.text.strip():
                    content.append(p.text.strip())
        elif element.tag.endswith('tbl'):
            for t in document.tables:
                if t._element == element:
                    for row in t.rows:
                        row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_data:
                            content.append(" | ".join(row_data))
                            
    if not content:
        # Fallback if sequential iteration fails
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text.strip())
        for table in document.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    content.append(" | ".join(row_data))

    return "\n".join(content).strip()

def extract_from_bytes(file_bytes: bytes, filename: str) -> str:
    """Extract and clean text from bytes based on the original filename extension."""
    path = Path(filename)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        print(f"Received file: {filename}\nDetected type: application/pdf\nSelected extractor: PDF")
        text = extract_pdf_bytes(file_bytes)
    elif suffix == ".docx":
        print(f"Received file: {filename}\nDetected type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\nSelected extractor: DOCX")
        text = extract_docx_bytes(file_bytes)
    elif suffix == ".txt":
        print(f"Received file: {filename}\nDetected type: text/plain\nSelected extractor: TXT")
        text = file_bytes.decode("utf-8")
    else:
        raise ValueError(f"Unsupported file format. Please upload a PDF or DOCX file.")
        
    if not text or not text.strip():
        raise ValueError("Could not extract readable text from this resume. Please upload a clearer file.")

    return clean_text(text)

def extract_text(file_path: str) -> str:
    """Extract and clean text based on the file extension."""
    path = Path(file_path)
    file_bytes = path.read_bytes()
    return extract_from_bytes(file_bytes, path.name)